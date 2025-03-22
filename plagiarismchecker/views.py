from django.shortcuts import render,redirect
from django.contrib import messages
from django.http import HttpResponse
from plagiarismchecker.algorithm import main
from docx import *
from plagiarismchecker.algorithm import fileSimilarity
import PyPDF2 
import nltk
from nltk.corpus import wordnet
# Create your views here.
#side menu

#home
def mainhome(request):
    return render(request, 'pc/landingPage.html')
    # return render(request, 'pc/homePage.html') 
def home(request):
    return render(request ,'pc/homePage.html')
    
def paraphrase(request):
    return render(request, 'pc/phraphrase.html') 
#web search(Text)
def test(request):
    print("request is welcome test")
    print(request.POST['q'])  
    
    if request.POST['q']: 
        percent,link = main.findSimilarity(request.POST['q'])
        percent = round(percent,2)
    print("Output.....................!!!!!!!!",percent,link)
    return render(request, 'pc/textUpload.html',{'link': link, 'percent': percent , "text": request.POST['q']})

#web search file(.txt, .docx)
def filetest(request):
    value = ''    
    print(request.FILES['docfile'])
    if str(request.FILES['docfile']).endswith(".txt"):
        value = str(request.FILES['docfile'].read())

    elif str(request.FILES['docfile']).endswith(".docx"):
        document = Document(request.FILES['docfile'])
        for para in document.paragraphs:
            value += para.text

    elif str(request.FILES['docfile']).endswith(".pdf"):
        # creating a pdf file object 
        # pdfFileObj = open(request.FILES['docfile'], 'rb') 
        docfile = request.FILES['docfile']
        # creating a pdf reader object 
        pdfReader = PyPDF2.PdfReader(docfile) 

        # printing number of pages in pdf file 
        print(pdfReader.pages) 

        # creating a page object 
        pageObj = pdfReader.pages[0] 

        # extracting text from page 
        print(pageObj.extract_text()) 
        value = pageObj.extract_text()
        # closing the pdf file object 
        # pdfFileObj.close() 


    percent,link = main.findSimilarity(value)
    print("Output...................!!!!!!!!",percent,link)
    return render(request, 'pc/documentUpload.html',{'link': link, 'percent': percent})

#text compare
def fileCompare(request):
    return render(request, 'pc/doc_compare.html') 

#two text compare(Text)
def twofiletest1(request):
    print("Submiited text for 1st and 2nd")
    print(request.POST['q1'])
    print(request.POST['q2'])

    if request.POST['q1'] != '' and request.POST['q2'] != '': 
        print("Got both the texts")
        result = fileSimilarity.findFileSimilarity(request.POST['q1'],request.POST['q2'])
    result = round(result,2)    
    print("Output>>>>>>>>>>>>>>>>>>>>!!!!!!!!",result)
    return render(request, 'pc/comparetextCheck.html',{'result': result ,"text1": request.POST['q1'], "text2": request.POST['q2']})
    

#two text compare(.txt, .docx)
def twofilecompare1(request):
    file1 = request.FILES['docfile1']
    file2 = request.FILES['docfile2']

    ext1 = request.FILES['docfile1'].name.split('.')[-1].lower()
    ext2 = request.FILES['docfile2'].name.split('.')[-1].lower()

    # Check if both files have the same extension
    if ext1 != ext2:
        messages.error(request, "Error: Files must have the same extension.")
        return redirect(comparefilecheck)
  

    value1 = ''
    value2 = ''

    # Get uploaded files
    

    # Process TXT files
    if ext1 == "txt" and ext2=="txt":
        value1 = file1.read().decode('utf-8', errors='ignore')  # Decode to string
        value2 = file2.read().decode('utf-8', errors='ignore')

    # Process DOCX files
    elif ext1 == "docx" and ext2=="docx":
        doc1 = Document(file1)
        doc2 = Document(file2)
        value1 = "\n".join([para.text for para in doc1.paragraphs])
        value2 = "\n".join([para.text for para in doc2.paragraphs])

    # Process PDF files
    elif ext1 == "pdf" and ext2=="pdf":
        def extract_pdf_text(pdf_file):
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()

        value1 = extract_pdf_text(file1)
        value2 = extract_pdf_text(file2)

    # Run similarity check
    result = fileSimilarity.findFileSimilarity(value1, value2)

    print("Output..................!!!!!!!!", result)
    return render(request, 'pc/comparefilecheck.html', {'result': result})

# Function to paraphrase text
def paraphrase_with_synonyms(text):
    words = text.split()
    new_words = []

    for word in words:
        synonyms = wordnet.synsets(word)  # Find synonyms
        if synonyms:
            new_word = synonyms[0].lemmas()[0].name()  # Pick the first synonym
            new_words.append(new_word)
        else:
            new_words.append(word)

    return ' '.join(new_words)

# handling form submission for pharaphrasing
def paraphrase_view(request):
    if request.method == "POST":
        input_text = request.POST.get("para", "")
        processed_text =paraphrase_with_synonyms(input_text)  # Replace with actual paraphrasing logic
        return render(request, "pc/paraphrase.html", {"output": processed_text , "inputText":input_text})

    return render(request, "pc/paraphrase.html")

def documentUpload(request):
    return render(request, 'pc/documentUpload.html')

def textUpload(request):
    return render(request, 'pc/textUpload.html')

def comparetextCheck(request):
    return render(request, 'pc/comparetextCheck.html')


def comparefilecheck(request):
    return render(request, 'pc/comparefilecheck.html')

def helpus(request):
    return render(request, 'pc/help.html')

def contactus(request):
    return render(request, 'pc/contactus.html')
